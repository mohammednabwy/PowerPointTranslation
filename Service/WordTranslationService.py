import os
import docx
from docx import Document
from Service import Translator


def save(filename,document):
    document.save(filename)
    return


def translateWordFile(wordfilePath,target='ar'): 
    document=Document(wordfilePath)   
    for paragraph in document.paragraphs:
        if target=='ar':
            pPr=paragraph._p.get_or_add_pPr()
            p_bidi= docx.oxml.shared.OxmlElement('w:bidi')
            pPr.append(p_bidi)
        else:
            pPr=paragraph._p.get_or_add_pPr()
            for child in pPr.getchildren():
             if 'bidi' in  child.__repr__():                 
                 pPr.remove(child)
        for run in paragraph.runs:     
            try:
                cur_text = run.text
                if len(cur_text) < 1 : continue
                new_text = Translator.translate(cur_text,target) 
                run.text = new_text 
                if target=='ar':                               
                    run.font.rtl = True
                else:
                    run.font.rtl = False
            except:
                continue
    
    for table in document.tables:
        try:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:                  
                            paragraph.text = Translator.translate(paragraph.text,target) 
        except:
                continue
    
    #save 
    file_name=os.path.basename(wordfilePath)
    file_directory=os.path.dirname(wordfilePath) 
    translated_file_name='translated_' + target +"_" +  file_name.split('.')[0] + '.'  +file_name.split('.')[1]
    translated_file = os.path.join(file_directory, translated_file_name)
    save(translated_file,document)
    return translated_file

def main(): 
     wordfilePath='Files\\programmin_1.docx'   
     translateWordFile(wordfilePath,target='en')  
     return